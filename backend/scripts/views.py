# views.py
import logging
import os
import time
from datetime import datetime

from django.conf import settings
from django.shortcuts import get_object_or_404
from django_filters.rest_framework import DjangoFilterBackend
from drf_spectacular.utils import (
    OpenApiExample,
    OpenApiParameter,
    OpenApiResponse,
    OpenApiTypes,
    extend_schema,
)
from rest_framework import filters, generics, status
from rest_framework.decorators import api_view, parser_classes, permission_classes
from rest_framework.parsers import FormParser, JSONParser, MultiPartParser
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

from api.mixins import MethodSpecificThrottleMixin
from notifications.choices import NotificationType
from notifications.helpers import NotificationHelper
from payments.permissions import HasActiveSubscriptionPermission

from .filters import FullScriptFilter, ScriptOutlineFilter, generation_filters
from .models import FullScript, ScriptOutline, TemplateStyle, Tone
from .pagination import GenerationsLimitOffsetPagination
from .serializers import (
    ExportScriptResponseSerializer,
    FullScriptSerializer,
    GenerateOutlineRequestSerializer,
    GenerateOutlineResponseSerializer,
    GenerateScriptRequestSerializer,
    GenerateScriptResponseSerializer,
    ScriptGeneratorConfigResponseSerializer,
    ScriptOutlineSerializer,
    ScriptOutlineUpdateSerializer,
    StatusUpdateSerializer,
    TemplateStyleSerializer,
    ToneSerializer,
    UnifiedGenerationSerializer,
)
from .services.open_ai import OpenAIScriptService

logger = logging.getLogger(__name__)


@extend_schema(
    summary="Get script generator configuration",
    description="Retrieve all configuration data needed for the script generator UI",
    responses={
        200: OpenApiResponse(
            response=ScriptGeneratorConfigResponseSerializer,
            description="Configuration data retrieved successfully",
        )
    },
)
@api_view(["GET"])
def script_generator_config(request):
    tones = Tone.objects.all().order_by("name")
    template_styles = TemplateStyle.objects.all().order_by("name")
    return Response(
        {
            "tones": ToneSerializer(tones, many=True).data,
            "template_styles": TemplateStyleSerializer(template_styles, many=True).data,
            "length_range": {"min": 0, "max": 10000, "default": 500},
        }
    )


@extend_schema(
    summary="Generate script outline",
    description="""Generate a detailed script outline using multiple tones.

**Three input methods supported:**
1. **Text Description (JSON)**: Provide a text description in JSON format
2. **Image Upload (Form Data)**: Upload an image file using multipart/form-data
3. **Image URL**: Provide an image URL that will be analyzed

**Usage:**
- For text input: Use `Content-Type: application/json` with `description` field
- For image file: Use `Content-Type: multipart/form-data` with `image` file
- For image URL: Use either JSON or form data with `image_url` field
- Either `description`, `image` file, or `image_url` must be provided (only one)
- **When providing an image**: No need to provide title or description - they will be automatically generated from the image content using OpenAI Vision""",
    request={
        "multipart/form-data": GenerateOutlineRequestSerializer,
        "application/json": GenerateOutlineRequestSerializer,
    },
    responses={
        201: OpenApiResponse(
            response=GenerateOutlineResponseSerializer,
            description="Script outline generated successfully",
        ),
        400: OpenApiResponse(
            description="Invalid input data - either description or image must be provided"
        ),
        403: OpenApiResponse(description="Active subscription required"),
        500: OpenApiResponse(description="Outline generation or image analysis failed"),
    },
    examples=[
        OpenApiExample(
            "Text Description (JSON)",
            value={
                "description": "How to learn Python programming from scratch",
                "tones": [1],
                "template_style": 2,
                "min_length": 200,
                "max_length": 800,
                "title": "Python Learning Guide",
            },
            media_type="application/json",
        ),
        OpenApiExample(
            "Image Upload (Form Data)",
            value={
                "tones": [1, 3],
                "template_style": 1,
                "min_length": 300,
                "max_length": 1000,
                "image": "[Upload image file here]",
            },
            media_type="multipart/form-data",
        ),
        OpenApiExample(
            "Image URL (JSON)",
            value={
                "tones": [1, 2],
                "template_style": 2,
                "min_length": 200,
                "max_length": 800,
                "image_url": "https://example.com/path/to/image.jpg",
            },
            media_type="application/json",
        ),
    ],
)
@api_view(["POST"])
@parser_classes([MultiPartParser, FormParser, JSONParser])
@permission_classes([IsAuthenticated, HasActiveSubscriptionPermission])
def generate_script_outline(request):
    logger.info(
        f"[OUTLINE_GENERATION] Starting outline generation for user: {request.user.id}"
    )
    
    serializer = GenerateOutlineRequestSerializer(data=request.data)
    if not serializer.is_valid():
        logger.warning(
            f"[OUTLINE_GENERATION] Validation failed for user {request.user.id}: {serializer.errors}"
        )
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    validated_data = serializer.validated_data
    description = validated_data.get("description", "")
    image = validated_data.get("image")
    image_url = validated_data.get("image_url", "")
    tone_ids = validated_data.get("tones", [])
    template_style_id = validated_data.get("template_style")
    min_length = validated_data.get("min_length", 100)
    max_length = validated_data.get("max_length", 1000)
    title = validated_data.get("title", "")
    
    logger.info(
        f"[OUTLINE_GENERATION] Parameters - User: {request.user.id}, Description length: {len(description)}, "
                f"Tone IDs: {tone_ids}, Template style: {template_style_id}, Length: {min_length}-{max_length}, "
        f"Has image: {bool(image)}, Has image URL: {bool(image_url)}, Has title: {bool(title)}"
    )

    if image or image_url:
        logger.info(
            f"[OUTLINE_GENERATION] Starting image analysis for user {request.user.id}"
        )
        try:
            if image:
                if hasattr(image, "read") and hasattr(image, "seek"):
                    logger.debug(
                        f"[OUTLINE_GENERATION] Analyzing uploaded image file for user {request.user.id}"
                    )
                    image_title, image_description = (
                        OpenAIScriptService.analyze_image_with_assistant(
                            image_file=image, user=request.user
                        )
                    )
                else:
                    logger.error(
                        f"[OUTLINE_GENERATION] Invalid image file provided by user {request.user.id}"
                    )
                    return Response(
                        {"error": "Invalid image file."},
                        status=status.HTTP_400_BAD_REQUEST,
                    )
            else:
                logger.debug(
                    f"[OUTLINE_GENERATION] Analyzing image URL for user {request.user.id}: {image_url}"
                )
                image_title, image_description = (
                    OpenAIScriptService.analyze_image_with_assistant(
                        image_url=image_url, user=request.user
                    )
                )

            title = image_title
            description = image_description
            logger.info(
                f"[OUTLINE_GENERATION] Image analysis completed for user {request.user.id} - "
                f"Generated title: '{title[:50]}...', Description length: {len(description)}"
            )

        except Exception as e:
            logger.error(
                f"[OUTLINE_GENERATION] Image analysis failed for user {request.user.id}: {str(e)}",
                exc_info=True,
            )
            return Response(
                {"error": "Failed to analyze the provided image."},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    try:
        # Handle optional tones - if no tones provided, use default "Informative" tone
        logger.debug(
            f"[OUTLINE_GENERATION] Processing tones for user {request.user.id}"
        )
        if tone_ids:
            tones = Tone.objects.filter(id__in=tone_ids)
            if len(tones) != len(tone_ids):
                logger.error(
                    f"[OUTLINE_GENERATION] Invalid tone IDs provided by user {request.user.id}: {tone_ids}"
                )
                return Response(
                    {"error": "One or more invalid tone IDs provided"},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            logger.info(
                f"[OUTLINE_GENERATION] Using {len(tones)} tones for user {request.user.id}: {[tone.name for tone in tones]}"
            )
        else:
            # Use default "Informative" tone if no tones provided
            default_tone = Tone.objects.filter(name__iexact="Informative").first()
            if not default_tone:
                # If "Informative" doesn't exist, get the first available tone
                default_tone = Tone.objects.first()
                if not default_tone:
                    logger.error(
                        f"[OUTLINE_GENERATION] No tones available in the system for user {request.user.id}"
                    )
                    return Response(
                        {"error": "No tones available in the system"},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    )
            tones = [default_tone]
            logger.info(
                f"[OUTLINE_GENERATION] Using default tone '{default_tone.name}' for user {request.user.id}"
            )

        template_style = None
        if template_style_id:
            logger.debug(
                f"[OUTLINE_GENERATION] Processing template style {template_style_id} for user {request.user.id}"
            )
            template_style = TemplateStyle.objects.get(id=template_style_id)
            # Use template style word ranges when template is selected
            min_length = template_style.min_length
            max_length = template_style.max_length
            logger.info(
                f"[OUTLINE_GENERATION] Template style '{template_style.name}' applied - Length range: {min_length}-{max_length}"
            )

        # Extract tone names for title generation
        outline_tones = [tone.name for tone in tones]
        logger.debug(
            f"[OUTLINE_GENERATION] Final tone names for user {request.user.id}: {outline_tones}"
        )

        script_data = {
            "description": description,
            "tones": outline_tones,
            "template_style": template_style.name if template_style else "medium",
            "template_style_id": template_style.id if template_style else None,
            "min_length": min_length,
            "max_length": max_length,
        }
        
        logger.info(
            f"[OUTLINE_GENERATION] Prepared script data for user {request.user.id}: "
                   f"Description length: {len(description)}, Tones: {outline_tones}, "
                   f"Template: {template_style.name if template_style else 'medium'}, "
            f"Length range: {min_length}-{max_length}"
        )

        # Use assistant for outline generation with knowledge base
        logger.info(
            f"[OUTLINE_GENERATION] Starting OpenAI outline generation for user {request.user.id}"
        )
        outline_text, outline_data, metadata = (
            OpenAIScriptService.generate_outline_with_assistant(
                script_data, user=request.user
            )
        )
        logger.info(
            f"[OUTLINE_GENERATION] OpenAI outline generation completed for user {request.user.id} - "
                   f"Tokens used: {metadata.get('tokens_used', 0)}, Generation time: {metadata.get('generation_time', 0):.2f}s, "
            f"Model: {metadata.get('model', 'unknown')}"
        )

        # Parse JSON response - OpenAI should always return valid JSON
        import json
        import re

        logger.debug(
            f"[OUTLINE_GENERATION] Processing outline response for user {request.user.id}"
        )
        
        # Check for empty response first
        if not outline_text or not outline_text.strip():
            logger.error(
                f"[OUTLINE_GENERATION] OpenAI returned empty response for user {request.user.id}"
            )
            raise ValueError(
                "OpenAI returned empty response - this indicates an API error or rate limiting"
            )
        
        try:
            # Parse JSON directly - API enforces JSON format
            logger.debug(
                f"[OUTLINE_GENERATION] Parsing JSON response length: {len(outline_text)}"
            )
            
            outline_json_data = json.loads(outline_text)
            if (
                not isinstance(outline_json_data, dict)
                or "sections" not in outline_json_data
            ):
                raise ValueError("Invalid JSON structure: missing 'sections' field")
            
            # Extract structured data from JSON response
            actual_outline_data = {
                "sections": outline_json_data.get("sections", []),
                "section_order": outline_json_data.get("section_order", []),
            }
            
            # Reconstruct outline text from sections to avoid redundancy
            sections = actual_outline_data.get("sections", [])
            if not sections:
                raise ValueError("Invalid JSON structure: missing 'sections' field")
            
            # Build outline text from sections
            outline_parts = []
            for section in sections:
                title = section.get("title", "")
                description = section.get("description", "")
                if title and description:
                    outline_parts.append(f"{title} - {description}")
            
            actual_outline_text = "\n\n".join(outline_parts)

            # Clean up any document references
            actual_outline_text = re.sub(r"■[^■]*?■", "", actual_outline_text)
            actual_outline_text = re.sub(r"[0-9]+:\d+†[^■]*?■", "", actual_outline_text)
            actual_outline_text = re.sub(r"\n\s*\n\s*\n", "\n\n", actual_outline_text)
            actual_outline_text = actual_outline_text.strip()
            
            logger.info(
                f"[OUTLINE_GENERATION] Parsed JSON outline for user {request.user.id} - "
                       f"Sections: {len(actual_outline_data.get('sections', []))}, "
                f"Outline text length: {len(actual_outline_text)}"
            )
                       
        except (json.JSONDecodeError, TypeError, ValueError) as e:
            logger.error(
                f"[OUTLINE_GENERATION] JSON parsing failed for user {request.user.id}: {str(e)}"
            )
            logger.error(
                f"[OUTLINE_GENERATION] Raw response length: {len(outline_text)}"
            )
            logger.error(
                f"[OUTLINE_GENERATION] Raw response (first 500 chars): {outline_text[:500]}..."
            )
            raise ValueError(f"OpenAI returned invalid JSON response: {str(e)}")

        # Generate title using assistant if not provided or if provided title is empty/whitespace
        if title and title.strip():
            outline_title = title.strip()
            logger.debug(
                f"[OUTLINE_GENERATION] Using provided title for user {request.user.id}: '{outline_title}'"
            )
        else:
            logger.info(
                f"[OUTLINE_GENERATION] Generating title for user {request.user.id} (provided title: '{title}')"
            )
            try:
                # Use the title generation assistant to create an engaging title
                generated_titles, title_metadata = OpenAIScriptService.generate_titles(
                    prompt=description,
                    title_count=1,  # We only need one title
                    tones=outline_tones,  # Use the same tones as the outline
                )
                if generated_titles and len(generated_titles) > 0:
                    outline_title = generated_titles[0].get(
                        "title", f"Outline: {description[:50]}"
                    )
                    logger.info(
                        f"[OUTLINE_GENERATION] Generated title for user {request.user.id}: '{outline_title}' "
                        f"(tokens: {title_metadata.get('tokens_used', 0)})"
                    )
                else:
                    outline_title = f"Outline: {description[:50]}"
                    logger.warning(
                        f"[OUTLINE_GENERATION] Title generation returned empty results for user {request.user.id}"
                    )
            except Exception as e:
                logger.warning(
                    f"[OUTLINE_GENERATION] Title generation failed for user {request.user.id}: {str(e)}, using fallback"
                )
                outline_title = f"Outline: {description[:50]}"
        # Save the template parameters in outline_data for later use
        outline_data_with_params = (
            actual_outline_data.copy() if actual_outline_data else {}
        )
        outline_data_with_params.update(
            {
                "template_style": template_style.name if template_style else "medium",
                "template_style_id": template_style.id if template_style else None,
                "min_length": min_length,
                "max_length": max_length,
            }
        )

        # Extract default section order from outline_data
        default_section_order = (
            actual_outline_data.get("section_order", []) if actual_outline_data else []
        )

        logger.info(
            f"[OUTLINE_GENERATION] Creating ScriptOutline object for user {request.user.id}"
        )
        outline = ScriptOutline.objects.create(
            user=request.user,
            title=outline_title,
            outline_text=actual_outline_text,
            outline_data=outline_data_with_params,  # Save with parameters
            section_order=default_section_order,  # Set default section order
            original_outline=actual_outline_text,
            status="generated",
            openai_model=metadata["model"],
            tokens_used=metadata["tokens_used"],
            generation_time=metadata["generation_time"],
        )
        outline.tones.set(tones)
        
        logger.info(
            f"[OUTLINE_GENERATION] Outline created successfully for user {request.user.id} - "
                   f"Outline ID: {outline.uuid}, Title: '{outline.title}', "
                   f"Total tokens: {metadata.get('tokens_used', 0)}, "
            f"Generation time: {metadata.get('generation_time', 0):.2f}s"
        )

        # Create notification for successful outline generation
        try:
            NotificationHelper.create_user_notification(
                user=request.user,
                title="Outline Generated Successfully!",
                message=f"Your outline '{outline_title}' has been generated successfully. You can now create a full script from it.",
                notification_type=NotificationType.SCRIPT,
                metadata={
                    "outline": {
                        "label": outline_title,
                        "link": f"/outlines/{outline.uuid}",
                    },
                },
            )
        except Exception as e:
            logger.warning(f"Failed to create outline success notification: {str(e)}")

        serializer = ScriptOutlineSerializer(outline)
        logger.info(
            f"[OUTLINE_GENERATION] Outline generation completed successfully for user {request.user.id}"
        )
        return Response(
            {
                "outline": serializer.data,
                "message": "Script outline generated successfully!",
            },
            status=status.HTTP_201_CREATED,
        )

    except TemplateStyle.DoesNotExist:
        logger.error(
            f"[OUTLINE_GENERATION] Invalid template style {template_style_id} for user {request.user.id}"
        )
        return Response(
            {"error": "Invalid template style selected"},
            status=status.HTTP_400_BAD_REQUEST,
        )
    except Exception as e:
        logger.error(f"Outline generation failed: {str(e)}")

        # Create notification for failed outline generation
        try:
            NotificationHelper.create_user_notification(
                user=request.user,
                title="Outline Generation Failed",
                message="We encountered an issue while generating your outline. Please try again or contact support if the problem persists.",
                notification_type=NotificationType.SCRIPT,
                metadata={
                    "outline": {
                        "label": "Generation Failed",
                        "link": "/dashboard",
                    },
                },
            )
        except Exception as notification_error:
            logger.warning(
                f"Failed to create outline failure notification: {str(notification_error)}"
            )

        logger.error(
            f"[OUTLINE_GENERATION] Outline generation failed for user {request.user.id}: {str(e)}",
            exc_info=True,
        )
        return Response(
            {"error": "Failed to generate outline."},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@extend_schema(
    summary="Recreate script outline",
    description="Recreate an existing script outline using the same parameters as the original. This will generate a new outline with the same tones, template style, and length parameters, but with fresh AI-generated content.",
    request=None,  # No request body needed
    responses={
        201: OpenApiResponse(
            response=GenerateOutlineResponseSerializer,
            description="Script outline recreated successfully",
        ),
        404: OpenApiResponse(description="Original outline not found or access denied"),
        403: OpenApiResponse(description="Active subscription required"),
        500: OpenApiResponse(description="Outline recreation failed"),
    },
)
@api_view(["POST"])
@permission_classes([IsAuthenticated, HasActiveSubscriptionPermission])
def recreate_script_outline(request, uuid):
    """
    Recreate a script outline using the same parameters as an existing outline.

    This endpoint takes the UUID of an existing outline and generates a new outline
    using the same parameters (tones, template style, description, etc.) but with
    fresh AI-generated content.
    """
    # Get the original outline, ensuring it belongs to the requesting user
    try:
        original_outline = ScriptOutline.objects.get(uuid=uuid, user=request.user)
    except ScriptOutline.DoesNotExist:
        return Response(
            {"error": "Outline not found or access denied"},
            status=status.HTTP_404_NOT_FOUND,
        )

    try:
        # Extract parameters from the original outline
        # We'll need to reconstruct the original parameters used for generation

        # Get tones from the original outline
        tones = list(original_outline.tones.all())
        if not tones:
            # Use default "Informative" tone if original outline has no tones
            default_tone = Tone.objects.filter(name__iexact="Informative").first()
            if not default_tone:
                # If "Informative" doesn't exist, get the first available tone
                default_tone = Tone.objects.first()
                if not default_tone:
                    return Response(
                        {"error": "No tones available in the system"},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    )
            tones = [default_tone]

        # Try to extract template style from the script if it exists
        template_style = None
        if original_outline.script and original_outline.script.template_style:
            template_style = original_outline.script.template_style

        # Use default length parameters (we could store these in outline_data if needed)
        min_length = 100
        max_length = 1000

        # Try to extract length parameters from outline_data if available
        if original_outline.outline_data:
            min_length = original_outline.outline_data.get("min_length", 100)
            max_length = original_outline.outline_data.get("max_length", 1000)

        # If template style is available, use its word ranges
        if template_style:
            min_length = template_style.min_length
            max_length = template_style.max_length

        # Extract description from the original outline
        # This could be stored in outline_data or we might need to use the title
        description = (
            original_outline.outline_data.get("description", "")
            if original_outline.outline_data
            else ""
        )
        if not description:
            # Fallback to using the title or outline text as description
            description = (
                original_outline.title
                or f"Recreate outline based on: {original_outline.outline_text[:100]}..."
            )

        # Prepare script data for OpenAI service
        script_data = {
            "description": description,
            "tones": [tone.name for tone in tones],
            "template_style": template_style.name if template_style else "medium",
            "min_length": min_length,
            "max_length": max_length,
        }

        # Generate new outline using assistant with knowledge base
        outline_text, outline_data, metadata = (
            OpenAIScriptService.generate_outline_with_assistant(
                script_data, user=request.user
            )
        )

        # Extract actual outline data from JSON response
        import json
        import re

        try:
            outline_json_data = json.loads(outline_text)
            if isinstance(outline_json_data, dict) and "sections" in outline_json_data:
                # Extract structured data from JSON response
                actual_outline_data = {
                    "sections": outline_json_data.get("sections", []),
                    "section_order": outline_json_data.get("section_order", []),
                    "outline_text": outline_json_data.get("outline_text", ""),
                }
                # Use the structured outline_text from JSON, not the raw JSON
                actual_outline_text = outline_json_data.get(
                    "outline_text", outline_text
                )

                # Clean up any document references
                actual_outline_text = re.sub(r"■[^■]*?■", "", actual_outline_text)
                actual_outline_text = re.sub(
                    r"[0-9]+:\d+†[^■]*?■", "", actual_outline_text
                )
                actual_outline_text = re.sub(
                    r"\n\s*\n\s*\n", "\n\n", actual_outline_text
                )
                actual_outline_text = actual_outline_text.strip()
            else:
                # Fallback to original data
                actual_outline_data = outline_data or {}
                actual_outline_text = outline_text
        except (json.JSONDecodeError, TypeError):
            # Fallback to original data
            actual_outline_data = outline_data or {}
            actual_outline_text = outline_text

        # Create new outline with "Recreated" prefix, but generate a fresh title
        if original_outline.title and original_outline.title.strip():
            # Generate a new title based on the original outline's content
            try:
                # Use the title generation assistant to create an engaging title
                generated_titles, title_metadata = OpenAIScriptService.generate_titles(
                    prompt=original_outline.description,  # Use the original description
                    title_count=1,  # We only need one title
                    tones=tones,  # Use the same tones as the new outline
                )
                if generated_titles and len(generated_titles) > 0:
                    new_title = f"Recreated: {generated_titles[0].get('title', original_outline.title)}"
                else:
                    new_title = f"Recreated: {original_outline.title}"
            except Exception as e:
                logger.warning(
                    f"Title generation failed during recreation: {str(e)}, using fallback"
                )
                new_title = f"Recreated: {original_outline.title}"
        else:
            new_title = f"Recreated outline from {original_outline.uuid}"

        # Extract default section order from outline_data
        default_section_order = (
            actual_outline_data.get("section_order", []) if actual_outline_data else []
        )

        new_outline = ScriptOutline.objects.create(
            user=request.user,
            title=new_title,
            outline_text=actual_outline_text,
            outline_data=actual_outline_data,
            section_order=default_section_order,  # Set default section order
            original_outline=actual_outline_text,
            status="generated",
            openai_model=metadata["model"],
            tokens_used=metadata["tokens_used"],
            generation_time=metadata["generation_time"],
        )

        # Set the same tones as the original
        new_outline.tones.set(tones)

        # Create notification for successful outline recreation
        try:
            NotificationHelper.create_user_notification(
                user=request.user,
                title="Outline Recreated Successfully!",
                message=f"Your outline '{new_title}' has been recreated successfully. You can now create a full script from it.",
                notification_type=NotificationType.SCRIPT,
                metadata={
                    "outline": {
                        "label": new_title,
                        "link": f"/outlines/{new_outline.uuid}",
                    },
                },
            )
        except Exception as e:
            logger.warning(
                f"Failed to create outline recreation success notification: {str(e)}"
            )

        serializer = ScriptOutlineSerializer(new_outline)
        return Response(
            {
                "outline": serializer.data,
                "message": f"Script outline recreated successfully from original outline {original_outline.uuid}!",
                "original_outline_uuid": str(original_outline.uuid),
            },
            status=status.HTTP_201_CREATED,
        )

    except Exception as e:
        logger.error(f"Outline recreation failed: {str(e)}")

        # Create notification for failed outline recreation
        try:
            NotificationHelper.create_user_notification(
                user=request.user,
                title="Outline Recreation Failed",
                message="We encountered an issue while recreating your outline. Please try again or contact support if the problem persists.",
                notification_type=NotificationType.SCRIPT,
                metadata={
                    "outline": {
                        "label": "Recreation Failed",
                        "link": "/dashboard",
                    },
                },
            )
        except Exception as notification_error:
            logger.warning(
                f"Failed to create outline recreation failure notification: {str(notification_error)}"
            )

        return Response(
            {"error": "Failed to recreate outline. Please try again."},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@extend_schema(
    summary="Get script outline",
    description="Retrieve a specific script outline by its UUID",
    responses={
        200: OpenApiResponse(description="Script outline retrieved successfully"),
        404: OpenApiResponse(description="Script outline not found"),
    },
)
class ScriptOutlineDetailView(
    MethodSpecificThrottleMixin, generics.RetrieveUpdateDestroyAPIView
):
    """
    Get and update script outline
    """

    serializer_class = ScriptOutlineSerializer
    lookup_field = "uuid"
    permission_classes = [IsAuthenticated, HasActiveSubscriptionPermission]

    def get_queryset(self):
        if getattr(self, "swagger_fake_view", False):
            return ScriptOutline.objects.none()
        return ScriptOutline.objects.filter(user=self.request.user).prefetch_related(
            "tones"
        )

    def get_serializer_class(self):
        if self.request.method == "PATCH":
            return ScriptOutlineUpdateSerializer
        return ScriptOutlineSerializer

    @extend_schema(
        summary="Get script outline",
        description="Retrieve a specific script outline by its UUID",
        responses={
            200: OpenApiResponse(
                response=ScriptOutlineSerializer,
                description="Script outline retrieved successfully",
            ),
            404: OpenApiResponse(description="Script outline not found"),
        },
    )
    def get(self, request, *args, **kwargs):
        return super().get(request, *args, **kwargs)

    @extend_schema(
        summary="Update script outline",
        description="Update a specific script outline by its UUID",
        request=ScriptOutlineUpdateSerializer,
        responses={
            200: OpenApiResponse(
                response=ScriptOutlineSerializer,
                description="Script outline updated successfully",
            ),
            404: OpenApiResponse(description="Script outline not found"),
            400: OpenApiResponse(description="Invalid input data"),
        },
    )
    def patch(self, request, *args, **kwargs):
        return super().patch(request, *args, **kwargs)

    @extend_schema(
        summary="Delete script outline",
        description="Delete a specific script outline by its UUID",
        responses={
            200: OpenApiResponse(description="Script outline deleted successfully"),
            404: OpenApiResponse(description="Script outline not found"),
        },
    )
    def delete(self, request, *args, **kwargs):
        instance = self.get_object()
        instance.delete()
        return Response(
            {"message": "Script outline deleted successfully"},
            status=status.HTTP_200_OK,
        )


@extend_schema(
    summary="Generate full script",
    description="Generate a complete script from an existing outline. The script will use the tones from the outline, or fallback to the provided tone if the outline has no tones.",
    request=GenerateScriptRequestSerializer,
    responses={
        201: OpenApiResponse(
            response=GenerateScriptResponseSerializer,
            description="Full script generated successfully",
        ),
        404: OpenApiResponse(description="Script outline not found"),
        500: OpenApiResponse(description="Script generation failed"),
    },
    examples=[
        OpenApiExample(
            "Full Script Example",
            value={
                "tone": "Informative",
                "template_style": "medium",
                "min_length": 1500,
                "max_length": 3000,
                "title": "Complete Python Tutorial",
            },
        )
    ],
)
@api_view(["POST"])
@permission_classes([IsAuthenticated, HasActiveSubscriptionPermission])
def generate_full_script(request, uuid):
    """
    Generate full script from outline
    """
    logger.info(
        f"[SCRIPT_GENERATION] Starting script generation for user: {request.user.id}, outline: {uuid}"
    )
    outline = get_object_or_404(ScriptOutline, uuid=uuid, user=request.user)
    logger.info(
        f"[SCRIPT_GENERATION] Found outline '{outline.title}' for user {request.user.id}"
    )

    try:
        # Get tones from the outline (already stored)
        outline_tones = [tone.name for tone in outline.tones.all()]
        if not outline_tones:
            outline_tones = ["Informative"]
        logger.info(
            f"[SCRIPT_GENERATION] Using tones for user {request.user.id}: {outline_tones}"
        )

        # Read template style and word ranges from outline_data (stored during generation)
        template_style_name = outline.outline_data.get("template_style", "medium")
        template_style_id = outline.outline_data.get("template_style_id")
        min_length = outline.outline_data.get("min_length", 1000)
        max_length = outline.outline_data.get("max_length", 5000)
        logger.info(
            f"[SCRIPT_GENERATION] Template parameters from outline for user {request.user.id}: "
            f"style='{template_style_name}', length={min_length}-{max_length}"
        )

        # Optional: Allow request data to override (for flexibility)
        if request.data.get("template_style_id"):
            logger.debug(
                f"[SCRIPT_GENERATION] Overriding template style for user {request.user.id}: {request.data['template_style_id']}"
            )
            template_style = TemplateStyle.objects.get(
                id=request.data["template_style_id"]
            )
            min_length = template_style.min_length
            max_length = template_style.max_length
            template_style_name = template_style.name
            logger.info(
                f"[SCRIPT_GENERATION] Template style overridden for user {request.user.id}: "
                f"'{template_style_name}', length={min_length}-{max_length}"
            )

        script_data = {
            "tones": outline_tones,
            "template_style": template_style_name,
            "min_length": min_length,
            "max_length": max_length,
        }
        logger.info(
            f"[SCRIPT_GENERATION] Prepared script data for user {request.user.id}: {script_data}"
        )

        # Prepare structured outline data for script generation
        import json

        logger.debug(
            f"[SCRIPT_GENERATION] Preparing outline data for user {request.user.id}"
        )
        if outline.outline_data and outline.outline_data.get("sections"):
            # Use structured outline data with section_order
            structured_outline = {
                "sections": outline.outline_data.get("sections", []),
                "section_order": outline.outline_data.get("section_order", []),
                "outline_text": outline.outline_text,
            }
            outline_for_script = json.dumps(structured_outline)
            logger.info(
                f"[SCRIPT_GENERATION] Using structured outline for user {request.user.id} - "
                       f"Sections: {len(outline.outline_data.get('sections', []))}, "
                f"Outline text length: {len(outline.outline_text)}"
            )
        else:
            # Fallback to plain text
            outline_for_script = outline.outline_text
            logger.info(
                f"[SCRIPT_GENERATION] Using plain text outline for user {request.user.id} - "
                f"Length: {len(outline_for_script)}"
            )

        # Special handling for Flexible Outline template (ID: 4) - return outline as script
        if template_style_id == 4:
            logger.info(
                f"[SCRIPT_GENERATION] Flexible outline template detected for user {request.user.id} - "
                f"returning outline as script content"
            )
            
            # For Flexible templates, the outline IS the script
            script_content = outline.outline_text
            
            # Create minimal sections structure for Flexible template
            sections = [{
                "title": "Outline",
                "content": script_content,
                "start_time": "00:00",
                "end_time": "00:01",
                "word_count": len(script_content.split()),
                "section_type": "outline"
            }]
            
            # Create metadata for Flexible template
            metadata = {
                "tokens_used": 0,
                "generation_time": 0.0,
                "model": "flexible_outline",
                "assistant_id": "flexible-outline",
                "vector_store_id": "none",
                "thread_id": f"flexible_{int(time.time())}",
                "run_id": f"flexible_run_{int(time.time())}",
                "file_search_used": False,
                "word_count": len(script_content.split()),
                "length_valid": True,
                "strategy_used": "flexible_outline",
                "sections_generated": 1,
                "template_style": template_style_name
            }
            
            script_response = {
                "full_text": script_content,
                "sections": sections,
                "metadata": metadata
            }
        else:
            # Generate full script using new word count strategy for other templates
            logger.info(
                f"[SCRIPT_GENERATION] Starting section-based script generation for user {request.user.id}"
            )
            script_response = OpenAIScriptService.generate_script_with_word_count_strategy(
                outline_for_script, script_data, user=request.user
            )

        # Extract components from JSON response
        script_content = script_response["full_text"]
        sections = script_response["sections"]
        metadata = script_response["metadata"]
        logger.info(
            f"[SCRIPT_GENERATION] OpenAI script generation completed for user {request.user.id} - "
            f"Tokens used: {metadata.get('tokens_used', 0)}, Generation time: {metadata.get('generation_time', 0):.2f}s, "
            f"Model: {metadata.get('model', 'unknown')}"
        )

        # Process script response - new word count strategy returns plain text
        import re

        logger.debug(
            f"[SCRIPT_GENERATION] Processing script response for user {request.user.id}"
        )

        # Check for empty response first
        if not script_content or not script_content.strip():
            logger.error(
                f"[SCRIPT_GENERATION] OpenAI returned empty response for user {request.user.id}"
            )
            raise ValueError(
                "OpenAI returned empty response - this indicates an API error or rate limiting"
            )

        # The new word count strategy returns plain text content directly
        actual_script_text = script_content
        logger.info(
            f"[SCRIPT_GENERATION] Using plain text script for user {request.user.id} - "
            f"Script text length: {len(actual_script_text)}"
        )

        # Note: Section order validation is handled within the word count strategy method
        # The sections list is already provided by the strategy method

        # Clean up any document references that might have slipped through
        # Remove patterns like ■3:11†YOUTUBE STORYTELLING STRATEGY■
        logger.debug(
            f"[SCRIPT_GENERATION] Cleaning script text for user {request.user.id}"
        )
        actual_script_text = re.sub(r"■[^■]*?■", "", actual_script_text)
        actual_script_text = re.sub(r"[0-9]+:\d+†[^■]*?■", "", actual_script_text)
        # Clean up any extra whitespace
        actual_script_text = re.sub(r"\n\s*\n\s*\n", "\n\n", actual_script_text)
        actual_script_text = actual_script_text.strip()
        logger.info(
            f"[SCRIPT_GENERATION] Final script text length for user {request.user.id}: {len(actual_script_text)}"
        )

        # Create FullScript
        logger.info(
            f"[SCRIPT_GENERATION] Creating FullScript object for user {request.user.id}"
        )
        script_title = request.data.get("title", outline.title)
        full_script = FullScript.objects.create(
            user=request.user,
            outline=outline,
            title=script_title,  # Use the outline title directly
            content=actual_script_text,
            sections=sections,
            status="generated",
            openai_model=metadata["model"],
            tokens_used=metadata["tokens_used"],
            generation_time=metadata["generation_time"],
        )

        # Update outline status
        outline.status = "saved"
        outline.save()
        logger.info(
            f"[SCRIPT_GENERATION] Updated outline status to 'saved' for user {request.user.id}"
        )

        logger.info(
            f"[SCRIPT_GENERATION] Script created successfully for user {request.user.id} - "
                   f"Script ID: {full_script.uuid}, Title: '{script_title}', "
                   f"Content length: {len(actual_script_text)}, "
                   f"Total tokens: {metadata.get('tokens_used', 0)}, "
            f"Generation time: {metadata.get('generation_time', 0):.2f}s"
        )

        # Create notification for successful script generation
        try:
            NotificationHelper.create_user_notification(
                user=request.user,
                title="Script Generated Successfully!",
                message=f"Your script '{full_script.title}' has been generated successfully. You can now edit, approve, or export it.",
                notification_type=NotificationType.SCRIPT,
                metadata={
                    "script": {
                        "label": full_script.title,
                        "link": f"/scripts/{full_script.uuid}",
                    },
                },
            )
        except Exception as e:
            logger.warning(f"Failed to create script success notification: {str(e)}")

        serializer = FullScriptSerializer(full_script)
        logger.info(
            f"[SCRIPT_GENERATION] Script generation completed successfully for user {request.user.id}"
        )
        return Response(
            {
                "script": serializer.data,
                "message": "Full script generated successfully!",
            },
            status=status.HTTP_201_CREATED,
        )

    except Exception as e:
        logger.error(f"Script generation failed: {str(e)}")

        # Create notification for failed script generation
        try:
            NotificationHelper.create_user_notification(
                user=request.user,
                title="Script Generation Failed",
                message="We encountered an issue while generating your script. Please try again or contact support if the problem persists.",
                notification_type=NotificationType.SCRIPT,
                metadata={
                    "script": {
                        "label": "Generation Failed",
                        "link": "/dashboard",
                    },
                },
            )
        except Exception as notification_error:
            logger.warning(
                f"Failed to create script failure notification: {str(notification_error)}"
            )

        logger.error(
            f"[SCRIPT_GENERATION] Script generation failed for user {request.user.id}: {str(e)}",
            exc_info=True,
        )
        return Response(
            {"error": "Failed to generate script. Please try again."},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


class FullScriptDetailView(
    MethodSpecificThrottleMixin, generics.RetrieveUpdateDestroyAPIView
):
    """
    Full CRUD for full scripts
    """

    serializer_class = FullScriptSerializer
    lookup_field = "uuid"
    permission_classes = [IsAuthenticated, HasActiveSubscriptionPermission]

    def get_queryset(self):
        if getattr(self, "swagger_fake_view", False):
            return FullScript.objects.none()
        return FullScript.objects.filter(user=self.request.user).select_related(
            "outline"
        )

    @extend_schema(
        summary="Get full script",
        description="Retrieve a specific full script by its UUID",
        responses={
            200: OpenApiResponse(
                response=FullScriptSerializer,
                description="Full script retrieved successfully",
            ),
            404: OpenApiResponse(description="Full script not found"),
        },
    )
    def get(self, request, *args, **kwargs):
        return super().get(request, *args, **kwargs)

    @extend_schema(
        summary="Update full script",
        description="Update a specific full script by its UUID",
        request=FullScriptSerializer,
        responses={
            200: OpenApiResponse(
                response=FullScriptSerializer,
                description="Full script updated successfully",
            ),
            404: OpenApiResponse(description="Full script not found"),
            400: OpenApiResponse(description="Invalid input data"),
        },
    )
    def patch(self, request, *args, **kwargs):
        return super().patch(request, *args, **kwargs)

    @extend_schema(
        summary="Delete full script",
        description="Delete a specific full script by its UUID",
        responses={
            204: OpenApiResponse(description="Full script deleted successfully"),
            404: OpenApiResponse(description="Full script not found"),
        },
    )
    def delete(self, request, *args, **kwargs):
        return super().delete(request, *args, **kwargs)


@extend_schema(
    summary="List all script outlines",
    description="Retrieve a paginated list of all script outlines with filtering options",
    parameters=[
        OpenApiParameter(
            name="search",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Search by title or content",
            required=False,
        ),
        OpenApiParameter(
            name="status",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Filter by status (generating, generated, edited, approved, script_generated)",
            required=False,
        ),
        OpenApiParameter(
            name="filter_type",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Filter by type: all, outline_drafts, script_drafts, saved",
            required=False,
        ),
        OpenApiParameter(
            name="ordering",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Order by field (created, modified, title). Use - for descending (e.g., -created)",
            required=False,
        ),
    ],
    responses={
        200: OpenApiResponse(description="Script outlines retrieved successfully")
    },
)
class ScriptOutlineListView(MethodSpecificThrottleMixin, generics.ListAPIView):
    """
    List all script outlines with filtering and search
    """

    serializer_class = ScriptOutlineSerializer
    filter_backends = [DjangoFilterBackend, filters.OrderingFilter]
    filterset_class = ScriptOutlineFilter
    ordering_fields = ["created", "modified", "title"]
    ordering = ["-created"]
    permission_classes = [IsAuthenticated, HasActiveSubscriptionPermission]

    def get_queryset(self):
        if getattr(self, "swagger_fake_view", False):
            return ScriptOutline.objects.none()
        return (
            ScriptOutline.objects.filter(
                user=self.request.user, full_scripts__isnull=True
            )
            .prefetch_related("tones")
            .order_by("-created")
        )


@extend_schema(
    summary="List all full scripts",
    description="Retrieve a paginated list of all full scripts with filtering options",
    parameters=[
        OpenApiParameter(
            name="search",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Search by title or content",
            required=False,
        ),
        OpenApiParameter(
            name="status",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Filter by status (generating, generated, edited, finalized, published)",
            required=False,
        ),
        OpenApiParameter(
            name="is_published",
            type=OpenApiTypes.BOOL,
            location=OpenApiParameter.QUERY,
            description="Filter by published status",
            required=False,
        ),
        OpenApiParameter(
            name="filter_type",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Filter by type: all, outline_drafts, script_drafts, saved",
            required=False,
        ),
        OpenApiParameter(
            name="word_count_min",
            type=OpenApiTypes.INT,
            location=OpenApiParameter.QUERY,
            description="Minimum word count",
            required=False,
        ),
        OpenApiParameter(
            name="word_count_max",
            type=OpenApiTypes.INT,
            location=OpenApiParameter.QUERY,
            description="Maximum word count",
            required=False,
        ),
        OpenApiParameter(
            name="duration_min",
            type=OpenApiTypes.FLOAT,
            location=OpenApiParameter.QUERY,
            description="Minimum estimated duration in minutes",
            required=False,
        ),
        OpenApiParameter(
            name="duration_max",
            type=OpenApiTypes.FLOAT,
            location=OpenApiParameter.QUERY,
            description="Maximum estimated duration in minutes",
            required=False,
        ),
        OpenApiParameter(
            name="ordering",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Order by field (created, modified, title, word_count, estimated_duration). Use - for descending (e.g., -created)",
            required=False,
        ),
    ],
    responses={
        200: OpenApiResponse(
            response=FullScriptSerializer(many=True),
            description="Full scripts retrieved successfully",
        )
    },
)
class FullScriptListView(MethodSpecificThrottleMixin, generics.ListAPIView):
    """
    List all full scripts with filtering and search
    """

    serializer_class = FullScriptSerializer
    filter_backends = [DjangoFilterBackend, filters.OrderingFilter]
    filterset_class = FullScriptFilter
    ordering_fields = [
        "created",
        "modified",
        "title",
        "word_count",
        "estimated_duration",
    ]
    ordering = ["-created"]
    permission_classes = [IsAuthenticated, HasActiveSubscriptionPermission]

    def get_queryset(self):
        if getattr(self, "swagger_fake_view", False):
            return FullScript.objects.none()
        return (
            FullScript.objects.filter(user=self.request.user)
            .select_related("outline")
            .order_by("-created")
        )


@extend_schema(
    summary="List all generations (unified)",
    description="""
    Retrieve a unified list of both script outlines and full scripts with comprehensive filtering and sorting options.

    **Filter Combinations:**
    - Use `type=script` to apply word_count and duration filters only to scripts
    - Use `type=outline` to get only outlines (word_count/duration filters ignored)
    - Combine multiple filters for precise results

    **Pagination:**
    - Use `limit` parameter to specify number of items per page (default: 20, max: 100)
    - Use `offset` parameter to specify starting position (default: 0)
    - Example: `?limit=10&offset=20` returns items 21-30

    **Examples:**
    - `?type=script&word_count_min=500&duration_max=10&limit=15` - Scripts with 500+ words, max 10 min, 15 per page
    - `?filter_type=saved&search=tutorial&limit=5` - Saved items containing 'tutorial', 5 per page
    - `?status=generated&ordering=-modified&limit=25&offset=50` - Generated items, newest first, page 3 (items 51-75)
    """,
    parameters=[
        OpenApiParameter(
            name="search",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Search by title or content across both outlines and scripts",
            required=False,
        ),
        OpenApiParameter(
            name="status",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Filter by status: draft, generated, saved",
            required=False,
        ),
        OpenApiParameter(
            name="filter_type",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Filter by type: all, outline_drafts, script_drafts, saved",
            required=False,
        ),
        OpenApiParameter(
            name="type",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Filter by generation type: outline, script",
            required=False,
        ),
        OpenApiParameter(
            name="ordering",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Order by field: created, modified, title, word_count, estimated_duration. Use - for descending (e.g., -created)",
            required=False,
        ),
        OpenApiParameter(
            name="limit",
            type=OpenApiTypes.INT,
            location=OpenApiParameter.QUERY,
            description="Number of items to return per page (default: 20, max: 100)",
            required=False,
        ),
        OpenApiParameter(
            name="offset",
            type=OpenApiTypes.INT,
            location=OpenApiParameter.QUERY,
            description="Number of items to skip from the beginning (default: 0)",
            required=False,
        ),
        OpenApiParameter(
            name="word_count_min",
            type=OpenApiTypes.INT,
            location=OpenApiParameter.QUERY,
            description="Minimum word count for scripts (only applies to scripts, outlines are unaffected). Example: 500",
            required=False,
        ),
        OpenApiParameter(
            name="word_count_max",
            type=OpenApiTypes.INT,
            location=OpenApiParameter.QUERY,
            description="Maximum word count for scripts (only applies to scripts, outlines are unaffected). Example: 2000",
            required=False,
        ),
        OpenApiParameter(
            name="duration_min",
            type=OpenApiTypes.FLOAT,
            location=OpenApiParameter.QUERY,
            description="Minimum estimated video duration in minutes (only applies to scripts, outlines are unaffected). Example: 3.5",
            required=False,
        ),
        OpenApiParameter(
            name="duration_max",
            type=OpenApiTypes.FLOAT,
            location=OpenApiParameter.QUERY,
            description="Maximum estimated video duration in minutes (only applies to scripts, outlines are unaffected). Example: 15.0",
            required=False,
        ),
    ],
    responses={
        200: OpenApiResponse(
            response=UnifiedGenerationSerializer(many=True),
            description="Generations retrieved successfully with pagination metadata",
        ),
        404: OpenApiResponse(description="No generations found for the given filters"),
        400: OpenApiResponse(description="Invalid filter parameters"),
    },
)
class GenerationsList(MethodSpecificThrottleMixin, generics.ListAPIView):
    """
    Unified listing API for both outlines and scripts with pagination.
    Handles filtering, sorting, and pagination for both outlines and scripts.
    """

    serializer_class = UnifiedGenerationSerializer
    pagination_class = GenerationsLimitOffsetPagination
    permission_classes = [IsAuthenticated, HasActiveSubscriptionPermission]

    def get_queryset(self):
        """
        Get filtered queryset based on request parameters.
        Returns a list of objects that will be serialized and paginated.
        """
        # Get filter parameters from request
        search = self.request.GET.get("search", "")
        status_filter = self.request.GET.get("status", "")
        filter_type = self.request.GET.get("filter_type", "all")
        type_filter = self.request.GET.get("type", "")
        ordering = self.request.GET.get("ordering", "-created")

        # Get numeric filters
        word_count_min = self.request.GET.get("word_count_min")
        word_count_max = self.request.GET.get("word_count_max")
        duration_min = self.request.GET.get("duration_min")
        duration_max = self.request.GET.get("duration_max")

        # Convert to integers if provided
        try:
            word_count_min = int(word_count_min) if word_count_min else None
        except (ValueError, TypeError):
            word_count_min = None

        try:
            word_count_max = int(word_count_max) if word_count_max else None
        except (ValueError, TypeError):
            word_count_max = None

        try:
            duration_min = float(duration_min) if duration_min else None
        except (ValueError, TypeError):
            duration_min = None

        try:
            duration_max = float(duration_max) if duration_max else None
        except (ValueError, TypeError):
            duration_max = None

        # Get filtered querysets using the existing filter function
        filtered_querysets = generation_filters(
            search,
            status_filter,
            filter_type,
            type_filter,
            ordering,
            user=self.request.user,
            word_count_min=word_count_min,
            word_count_max=word_count_max,
            duration_min=duration_min,
            duration_max=duration_max,
        )

        return filtered_querysets


@extend_schema(
    summary="Update outline status",
    description="Update the status of a script outline",
    request=StatusUpdateSerializer,
    responses={
        200: OpenApiResponse(
            response=ScriptOutlineSerializer,
            description="Outline status updated successfully",
        ),
        404: OpenApiResponse(description="Outline not found"),
        400: OpenApiResponse(description="Invalid status"),
    },
)
@api_view(["PATCH"])
@permission_classes([IsAuthenticated, HasActiveSubscriptionPermission])
def update_outline_status(request, outline_uuid):
    """
    Update the status of a script outline
    """
    outline = get_object_or_404(ScriptOutline, uuid=outline_uuid, user=request.user)

    serializer = StatusUpdateSerializer(data=request.data)
    if serializer.is_valid():
        outline.status = serializer.validated_data["status"]
        outline.save()

        response_serializer = ScriptOutlineSerializer(outline)
        return Response(response_serializer.data)

    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@extend_schema(
    summary="Update script status",
    description="Update the status of a full script",
    request=StatusUpdateSerializer,
    responses={
        200: OpenApiResponse(
            response=FullScriptSerializer,
            description="Script status updated successfully",
        ),
        404: OpenApiResponse(description="Script not found"),
        400: OpenApiResponse(description="Invalid status"),
    },
)
@api_view(["PATCH"])
@permission_classes([IsAuthenticated, HasActiveSubscriptionPermission])
def update_script_status(request, script_uuid):
    """
    Update the status of a full script
    """
    script = get_object_or_404(FullScript, uuid=script_uuid, user=request.user)

    serializer = StatusUpdateSerializer(data=request.data)
    if serializer.is_valid():
        script.status = serializer.validated_data["status"]
        script.save()

        response_serializer = FullScriptSerializer(script)
        return Response(response_serializer.data)

    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@extend_schema(
    summary="Test word count strategy",
    description="Test the new word count completion strategy with different template styles",
    responses={
        200: OpenApiResponse(description="Word count strategy test results"),
    },
)
@api_view(["GET"])
@permission_classes([IsAuthenticated, HasActiveSubscriptionPermission])
def test_word_count_strategy(request):
    """
    Test endpoint for the new word count strategy
    """
    from scripts.services.word_count_strategy import WordCountStrategy

    template_style = request.GET.get("template", "medium")

    try:
        strategy = WordCountStrategy(template_style)
        targets = strategy.calculate_section_word_targets()

        # Test different section counts
        test_results = {
            "template_style": template_style,
            "default_targets": targets,
            "custom_sections_test": {},
        }

        # Test with different section counts
        for sections in [4, 6, 8]:
            custom_targets = strategy.calculate_section_word_targets(sections)
            test_results["custom_sections_test"][sections] = {
                "total_target": custom_targets["total_target"],
                "intro_words": custom_targets["intro"],
                "main_words": custom_targets["main_sections"],
                "conclusion_words": custom_targets["conclusion"],
            }

        return Response(
            {
                "message": "Word count strategy test completed successfully",
                "results": test_results,
            }
        )

    except Exception as e:
        logger.error(f"Word count strategy test failed: {str(e)}")
        return Response(
            {"error": f"Test failed: {str(e)}"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@extend_schema(
    summary="Export script to file",
    description="Export a full script to various file formats (.txt, .pdf, .docx). The file will be downloaded with appropriate headers.",
    parameters=[
        OpenApiParameter(
            name="format",
            type=OpenApiTypes.STR,
            location=OpenApiParameter.QUERY,
            description="Export format: txt, pdf, or docx",
            required=True,
            enum=["txt", "pdf", "docx"],
        ),
    ],
    responses={
        200: OpenApiResponse(description="File exported successfully"),
        400: OpenApiResponse(description="Invalid format specified"),
        404: OpenApiResponse(description="Script not found or access denied"),
        403: OpenApiResponse(description="Active subscription required"),
    },
)
class ExportScriptView(APIView):
    """
    Export a full script to various file formats

    Supported formats:
    - txt: Plain text format
    - pdf: PDF format with proper formatting
    - docx: Microsoft Word document format
    """

    permission_classes = [IsAuthenticated]
    serializer_class = ExportScriptResponseSerializer

    def post(self, request, *args, **kwargs):
        # return Response({"message": "Export request received"}, status=status.HTTP_200_OK)
        """Handle GET request to export script in specified format"""
        # Debug logging
        logger.info(f"Export request received for UUID: {kwargs['uuid']}")
        logger.info(f"User: {request.user}")
        logger.info(f"Format: {request.data.get('format', 'not provided')}")

        # Get the format parameter
        export_format = request.data.get("format", "").lower()
        if export_format not in ["txt", "pdf", "docx"]:
            return Response(
                {"error": "Invalid format. Supported formats: txt, pdf, docx"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Get the script, ensuring it belongs to the requesting user
        try:
            script = FullScript.objects.get(uuid=kwargs["uuid"], user=request.user)
            logger.info(f"Script found: {script.title}")
        except FullScript.DoesNotExist:
            # Enhanced error logging
            logger.error(
                f"Script not found for UUID: {kwargs['uuid']}, User: {request.user}"
            )

            # Check if script exists for any user
            try:
                script_exists = FullScript.objects.get(uuid=kwargs["uuid"])
                logger.error(
                    f"Script exists but belongs to different user: {script_exists.user}"
                )
                return Response(
                    {
                        "error": "Script not found or access denied",
                        "debug": "Script exists but access denied",
                    },
                    status=status.HTTP_404_NOT_FOUND,
                )
            except FullScript.DoesNotExist:
                logger.error(f"Script with UUID {kwargs['uuid']} does not exist at all")
                return Response(
                    {
                        "error": "Script not found or access denied",
                        "debug": "Script does not exist",
                    },
                    status=status.HTTP_404_NOT_FOUND,
                )

        try:
            # Generate filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(
                c for c in script.title if c.isalnum() or c in (" ", "-", "_")
            ).rstrip()
            safe_title = safe_title.replace(" ", "_")[
                :50
            ]  # Limit length and replace spaces

            if export_format == "txt":
                result = _export_txt(script, safe_title, timestamp)
            elif export_format == "pdf":
                result = _export_pdf(script, safe_title, timestamp)
            elif export_format == "docx":
                result = _export_docx(script, safe_title, timestamp)
            result = request.build_absolute_uri(result["file_url"])
            return Response({"file_url": result}, status=status.HTTP_200_OK)

        except Exception as e:
            logger.error(f"Script export failed: {str(e)}")
            return Response(
                {"error": "Failed to export script. Please try again."},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )


def _export_txt(script, safe_title, timestamp):
    """Export script as plain text file and save to media folder"""
    import json

    content = f"Title: {script.title}\n"
    content += f"Created: {script.created.strftime('%Y-%m-%d %H:%M:%S')}\n"
    content += f"Word Count: {script.word_count}\n"
    content += f"Estimated Duration: {script.estimated_duration:.1f} minutes\n"
    content += "=" * 50 + "\n\n"

    # Content should now be clean script text (not JSON)
    # Try to parse as JSON first for backward compatibility
    try:
        script_data = json.loads(script.content)
        if isinstance(script_data, dict):
            # Handle legacy JSON format
            if "full_text" in script_data and "sections" in script_data:
                # Legacy JSON format: extract full_text
                content += script_data["full_text"]
            elif "script" in script_data:
                # Legacy format: nested script structure
                script_sections = script_data["script"]
                if isinstance(script_sections, dict) and "sections" in script_sections:
                    # Handle script sections
                    for section in script_sections["sections"]:
                        if "title" in section:
                            content += f"\n=== {section['title']} ===\n\n"
                        if "content" in section:
                            content += f"{section['content']}\n\n"
                elif isinstance(script_sections, list):
                    # Handle legacy list format
                    for section in script_sections:
                        if "section" in section:
                            content += f"\n=== {section['section']} ===\n\n"
                        if "content" in section and isinstance(
                            section["content"], list
                        ):
                            for item in section["content"]:
                                content += f"{item}\n\n"
                        elif "content" in section and isinstance(
                            section["content"], str
                        ):
                            content += f"{section['content']}\n\n"
                else:
                    # Fallback to original content if not in expected format
                    content += script.content
        else:
            # Fallback to original content if not in expected format
            content += script.content
    except (json.JSONDecodeError, KeyError, TypeError):
        # Content is now clean script text (new format)
        content += script.content

    filename = f"{safe_title}_{timestamp}.txt"
    filepath = os.path.join("exports", filename)  # subfolder in MEDIA_ROOT

    full_path = os.path.join(settings.MEDIA_ROOT, filepath)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)

    with open(full_path, "w", encoding="utf-8") as f:
        f.write(content)

    file_url = settings.MEDIA_URL + filepath
    return {"file_url": file_url}


def _export_pdf(script, safe_title, timestamp):
    import json
    from io import BytesIO

    from pypdf import PdfReader, PdfWriter
    from reportlab.lib.colors import darkblue
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    filename = f"{safe_title}_{timestamp}.pdf"
    filepath = os.path.join("exports", filename)
    full_path = os.path.join(settings.MEDIA_ROOT, filepath)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)

    # Create PDF content using ReportLab (for content generation)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
    )

    # Get default styles and create custom styles
    styles = getSampleStyleSheet()

    # Custom styles for better formatting
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=darkblue,
        fontName="Helvetica-Bold",
    )

    section_style = ParagraphStyle(
        "SectionHeader",
        parent=styles["Heading2"],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=16,
        textColor=darkblue,
        fontName="Helvetica-Bold",
    )

    content_style = ParagraphStyle(
        "CustomContent",
        parent=styles["Normal"],
        fontSize=11,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
        leftIndent=0,
        rightIndent=0,
        fontName="Helvetica",
    )

    # Build the story with better formatting
    story = []

    # Add title
    story.append(Paragraph(script.title, title_style))
    story.append(Spacer(1, 20))

    # Try to parse as JSON structure first
    try:
        script_data = json.loads(script.content)
        if isinstance(script_data, dict):
            # Handle new script schema format
            if "full_text" in script_data and "sections" in script_data:
                # New format: use full_text for main content
                content_lines = script_data["full_text"].split("\n")
                for line in content_lines:
                    clean_line = line.strip()
                    if clean_line:
                        story.append(Paragraph(clean_line, content_style))
                        story.append(Spacer(1, 4))
            elif "script" in script_data:
                # Legacy format: nested script structure
                script_sections = script_data["script"]
                if isinstance(script_sections, dict) and "sections" in script_sections:
                    # Handle new script sections format
                    for section in script_sections["sections"]:
                        if "title" in section:
                            story.append(Paragraph(section["title"], section_style))
                            story.append(Spacer(1, 8))
                        if "content" in section:
                            clean_content = section["content"].replace("//", "").strip()
                            if clean_content:
                                story.append(Paragraph(clean_content, content_style))
                                story.append(Spacer(1, 8))
                elif isinstance(script_sections, list):
                    # Handle legacy list format
                    for section in script_sections:
                        if "section" in section:
                            story.append(Paragraph(section["section"], section_style))
                            story.append(Spacer(1, 8))

                        if "content" in section and isinstance(
                            section["content"], list
                        ):
                            for item in section["content"]:
                                # Clean up the content for PDF formatting
                                clean_item = str(item).replace("//", "").strip()
                                if clean_item:
                                    # Check for special formatting
                                    if (
                                        clean_item.startswith("NARRATOR")
                                        or clean_item.startswith("CUT TO")
                                        or clean_item.startswith("VISUAL")
                                    ):
                                        # Format as stage directions
                                        story.append(
                                            Paragraph(
                                                f"<i>{clean_item}</i>", content_style
                                            )
                                        )
                                    else:
                                        story.append(
                                            Paragraph(clean_item, content_style)
                                        )
                                    story.append(Spacer(1, 4))
                        elif "content" in section and isinstance(
                            section["content"], str
                        ):
                            clean_content = section["content"].replace("//", "").strip()
                            if clean_content:
                                story.append(Paragraph(clean_content, content_style))
                                story.append(Spacer(1, 8))
        else:
            # Fallback to original content processing
            content_lines = script.content.split("\n")
            for line in content_lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                    continue

                # Handle different types of content
                if line.startswith("[") and line.endswith("]"):
                    story.append(Paragraph(line, section_style))
                elif line.startswith("**") and line.endswith("**"):
                    bold_style = ParagraphStyle(
                        "BoldText",
                        parent=styles["Normal"],
                        fontSize=11,
                        spaceAfter=8,
                        spaceBefore=6,
                        fontName="Helvetica-Bold",
                    )
                    story.append(Paragraph(line, bold_style))
                else:
                    story.append(Paragraph(line, content_style))
    except (json.JSONDecodeError, KeyError, TypeError):
        # If JSON parsing fails, use original content processing
        content_lines = script.content.split("\n")
        for line in content_lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue

            # Handle different types of content
            if line.startswith("[") and line.endswith("]"):
                story.append(Paragraph(line, section_style))
            elif line.startswith("**") and line.endswith("**"):
                bold_style = ParagraphStyle(
                    "BoldText",
                    parent=styles["Normal"],
                    fontSize=11,
                    spaceAfter=8,
                    spaceBefore=6,
                    fontName="Helvetica-Bold",
                )
                story.append(Paragraph(line, bold_style))
            else:
                story.append(Paragraph(line, content_style))

    doc.build(story)

    # Get the PDF content from buffer
    buffer.seek(0)
    pdf_content = buffer.getvalue()
    buffer.close()

    # Use pypdf to process and write the final PDF
    pdf_writer = PdfWriter()

    # Create a temporary PDF reader from the content
    temp_buffer = BytesIO(pdf_content)
    pdf_reader = PdfReader(temp_buffer)

    # Add all pages to the writer using pypdf
    for page in pdf_reader.pages:
        pdf_writer.add_page(page)

    # Write the final PDF using pypdf's write method
    with open(full_path, "wb") as output_file:
        pdf_writer.write(output_file)

    temp_buffer.close()

    file_url = settings.MEDIA_URL + filepath
    return {"file_url": file_url}


def _export_docx(script, safe_title, timestamp):
    import json

    from docx import Document

    filename = f"{safe_title}_{timestamp}.docx"
    filepath = os.path.join("exports", filename)
    full_path = os.path.join(settings.MEDIA_ROOT, filepath)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)

    doc = Document()

    # Add title
    doc.add_heading(script.title, 0)

    # Add metadata
    doc.add_paragraph(f"Created: {script.created.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Word Count: {script.word_count}")
    doc.add_paragraph(f"Estimated Duration: {script.estimated_duration:.1f} minutes")
    doc.add_paragraph("=" * 50)
    doc.add_paragraph()  # Add space

    # Try to parse as JSON structure first
    try:
        script_data = json.loads(script.content)
        if isinstance(script_data, dict):
            # Handle new script schema format
            if "full_text" in script_data and "sections" in script_data:
                # New format: use full_text for main content
                content_lines = script_data["full_text"].split("\n")
                for line in content_lines:
                    clean_line = line.strip()
                    if clean_line:
                        doc.add_paragraph(clean_line)
            elif "script" in script_data:
                # Legacy format: nested script structure
                script_sections = script_data["script"]
                if isinstance(script_sections, dict) and "sections" in script_sections:
                    # Handle new script sections format
                    for section in script_sections["sections"]:
                        if "title" in section:
                            # Add section heading
                            doc.add_heading(section["title"], level=1)
                        if "content" in section:
                            clean_content = section["content"].replace("//", "").strip()
                            if clean_content:
                                doc.add_paragraph(clean_content)
                elif isinstance(script_sections, list):
                    # Handle legacy list format
                    for section in script_sections:
                        if "section" in section:
                            # Add section heading
                            doc.add_heading(section["section"], level=1)

                        if "content" in section and isinstance(
                            section["content"], list
                        ):
                            for item in section["content"]:
                                # Clean up the content for DOCX formatting
                                clean_item = str(item).replace("//", "").strip()
                                if clean_item:
                                    # Check for special formatting
                                    if (
                                        clean_item.startswith("NARRATOR")
                                        or clean_item.startswith("CUT TO")
                                        or clean_item.startswith("VISUAL")
                                    ):
                                        # Format as italic stage directions
                                        p = doc.add_paragraph()
                                        p.add_run(clean_item).italic = True
                                    else:
                                        doc.add_paragraph(clean_item)
                        elif "content" in section and isinstance(
                            section["content"], str
                        ):
                            clean_content = section["content"].replace("//", "").strip()
                            if clean_content:
                                doc.add_paragraph(clean_content)
            else:
                # Fallback to original content if not in expected format
                doc.add_paragraph(script.content)
        else:
            # Fallback to original content if not in expected format
            doc.add_paragraph(script.content)
    except (json.JSONDecodeError, KeyError, TypeError):
        # If JSON parsing fails, use original content
        doc.add_paragraph(script.content)

    doc.save(full_path)
    # TODO: Add file to storage/DELETE AFTER 10 MINUTES
    file_url = settings.MEDIA_URL + filepath
    return {"file_url": file_url}
